"""
scraper_cenet.py
----------------
Descarga el contenido de las páginas de un curso Moodle (cenet.inet.edu.ar)
y lo guarda como archivo .docx con texto e imágenes.

Dependencias:
    pip install requests beautifulsoup4 python-docx lxml

Uso:
    python scraper_cenet.py
"""

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import getpass
import os
import re
import sys
import io
from urllib.parse import urljoin, urlparse

BASE_URL = "https://cenet.inet.edu.ar"
LOGIN_URL = f"{BASE_URL}/login/index.php"


# ─── Autenticación ────────────────────────────────────────────────────────────

def login(session: requests.Session, username: str, password: str) -> bool:
    """Realiza el login en el sitio Moodle y devuelve True si fue exitoso."""
    resp = session.get(LOGIN_URL, timeout=30)
    resp.raise_for_status()

    soup = BeautifulSoup(resp.text, "lxml")

    # Moodle usa un token CSRF llamado 'logintoken'
    token_input = soup.find("input", {"name": "logintoken"})
    logintoken = token_input["value"] if token_input else ""

    payload = {
        "username": username,
        "password": password,
        "logintoken": logintoken,
        "anchor": "",
    }

    resp = session.post(LOGIN_URL, data=payload, timeout=30)
    resp.raise_for_status()

    # Si el login fue exitoso, la URL ya no apunta al login
    if "login" in resp.url or "login" in resp.text[:2000].lower() and "loginerrormessage" in resp.text:
        return False
    return True


# ─── Scraping de contenido ────────────────────────────────────────────────────

def get_page_content(session: requests.Session, url: str):
    """
    Obtiene el título y el contenido principal de una página del curso.
    Devuelve (titulo, soup_del_contenido).
    """
    resp = session.get(url, timeout=30)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "lxml")

    # Título de la página
    titulo_tag = soup.find("h1") or soup.find("h2")
    titulo = titulo_tag.get_text(strip=True) if titulo_tag else "Sin título"

    # Contenido específico: el div con el material de la clase
    # Moodle lo pone en div.box.generalbox.center.clearfix (con posibles clases extra como py-3)
    contenido = (
        soup.find("div", class_=lambda c: c and all(cls in c for cls in ["box", "generalbox", "center", "clearfix"]))
        or soup.find("div", {"id": "region-main"})
        or soup.find("div", class_="region-main")
        or soup.find("div", {"role": "main"})
        or soup.find("main")
    )

    return titulo, contenido


def download_image(session: requests.Session, img_url: str) -> bytes | None:
    """Descarga una imagen y devuelve sus bytes, o None si falla."""
    try:
        resp = session.get(img_url, timeout=20)
        content_type = resp.headers.get("Content-Type", "")
        if resp.status_code == 200 and ("image" in content_type or img_url.lower().endswith((".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg"))):
            return resp.content
        print(f"  [!] Imagen no accesible (HTTP {resp.status_code}): {img_url}")
    except Exception as e:
        print(f"  [!] No se pudo descargar imagen: {img_url} — {e}")
    return None


def extraer_src_imagen(nodo) -> str:
    """Extrae la URL real de una imagen probando todos los atributos posibles."""
    for attr in ('src', 'data-src', 'data-lazy-src', 'data-original',
                 'data-lazy', 'data-url', 'data-full-url', 'data-hi-res-src'):
        val = nodo.get(attr, '').strip()
        if val:  # incluye data URIs base64
            return val
    return ''


def extraer_url_video(nodo) -> str:
    """
    Dado un <iframe> o <video>, devuelve la URL pública del video.
    Convierte URLs de embed de YouTube a URLs watch normales.
    """
    src = nodo.get('src', '') or nodo.get('data-src', '')
    if not src and nodo.name == 'video':
        source = nodo.find('source')
        src = source.get('src', '') if source else ''
    if not src:
        return ''
    # YouTube embed -> watch URL
    m = re.search(r'youtube(?:-nocookie)?\.com/embed/([\w-]+)', src)
    if m:
        return f"https://www.youtube.com/watch?v={m.group(1)}"
    # youtu.be short
    m = re.search(r'youtu\.be/([\w-]+)', src)
    if m:
        return f"https://www.youtube.com/watch?v={m.group(1)}"
    # Otros reproductores: devolver la src tal cual
    return src


def agregar_marcador_video(doc: Document, url: str, titulo: str = ""):
    """Inserta un bloque visible con el link al video."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label = f"VIDEO SUGERIDO{': ' + titulo if titulo else ''}"
    run_label = p.add_run(label + "\n")
    run_label.bold = True
    run_label.font.size = Pt(13)
    run_label.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
    run_url = p.add_run(url)
    run_url.font.size = Pt(11)
    run_url.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
    run_url.underline = True
    doc.add_paragraph()


def agregar_marcador_imagen(doc: Document, src: str, alt: str = ""):
    """Inserta un marcador visual en rojo cuando una imagen no pudo descargarse."""
    texto = f"[ AQUÍ VA LA FOTO{': ' + alt if alt else ''} ]"
    parrafo = doc.add_paragraph()
    parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = parrafo.add_run(texto)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # rojo
    # URL de la imagen como referencia
    ref = doc.add_paragraph(f"({src})")
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref.runs[0].font.size = Pt(8)
    ref.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)


# ─── Helpers de formato ──────────────────────────────────────────────────────

def css_color_a_rgb(css: str):
    """Convierte un color CSS básico a RGBColor. Devuelve None si no reconoce."""
    css = css.strip().lower()
    m = re.match(r'#([0-9a-f]{2})([0-9a-f]{2})([0-9a-f]{2})', css)
    if m:
        return RGBColor(int(m.group(1), 16), int(m.group(2), 16), int(m.group(3), 16))
    m = re.match(r'#([0-9a-f])([0-9a-f])([0-9a-f])$', css)
    if m:
        return RGBColor(int(m.group(1) * 2, 16), int(m.group(2) * 2, 16), int(m.group(3) * 2, 16))
    m = re.match(r'rgb\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\)', css)
    if m:
        return RGBColor(int(m.group(1)), int(m.group(2)), int(m.group(3)))
    NOMBRES = {
        'red': (255, 0, 0), 'blue': (0, 0, 255), 'green': (0, 128, 0),
        'navy': (0, 0, 128), 'gray': (128, 128, 128), 'grey': (128, 128, 128),
        'black': (0, 0, 0), 'white': (255, 255, 255), 'orange': (255, 165, 0),
        'purple': (128, 0, 128), 'darkblue': (0, 0, 139), 'teal': (0, 128, 128),
    }
    if css in NOMBRES:
        return RGBColor(*NOMBRES[css])
    return None


def agregar_inline(nodo, parrafo, bold=False, italic=False, underline=False, color=None):
    """
    Recorre el contenido inline de un nodo y agrega runs con formato al párrafo.
    Preserva negrita, cursiva, subrayado y colores de texto.
    """
    from bs4 import NavigableString

    if isinstance(nodo, NavigableString):
        texto = str(nodo)
        # Ignorar saltos de línea puros pero conservar espacios
        if not texto or texto in ('\n', '\r\n', '\r'):
            return
        run = parrafo.add_run(texto)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if underline:
            run.underline = True
        if color:
            try:
                run.font.color.rgb = color
            except Exception:
                pass
        return

    nombre = nodo.name
    if not nombre or nombre in ('script', 'style', 'noscript'):
        return

    # Acumular formato de la etiqueta actual
    nb = bold or nombre in ('strong', 'b')
    ni = italic or nombre in ('em', 'i')
    nu = underline or nombre in ('u', 'ins')
    nc = color

    # Links → color azul
    if nombre == 'a':
        nc = nc or RGBColor(0x00, 0x56, 0xB3)

    # Leer color del atributo style
    style_attr = nodo.get('style', '')
    if style_attr:
        m = re.search(r'\bcolor\s*:\s*([^;]+)', style_attr, re.IGNORECASE)
        if m:
            c = css_color_a_rgb(m.group(1).strip())
            if c:
                nc = c

    for hijo in nodo.children:
        agregar_inline(hijo, parrafo, nb, ni, nu, nc)


def tiene_hijos_bloque(nodo) -> bool:
    """Devuelve True si el nodo contiene algún elemento de bloque o imagen en cualquier nivel."""
    BLOQUE_TAGS = {
        'p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'ul', 'ol', 'table', 'blockquote', 'pre', 'section',
        'article', 'figure', 'figcaption', 'img', 'picture',
    }
    # Primero revisamos hijos directos
    if any(getattr(h, 'name', None) in BLOQUE_TAGS for h in nodo.children):
        return True
    # Si no encontramos nada directo, buscamos imgs en cualquier nivel del subárbol
    return bool(nodo.find('img'))


# ─── Construcción del documento Word ─────────────────────────────────────────

def insertar_imagen(doc, session, src, alt, page_url):
    """Descarga e inserta una imagen en el doc, o pone marcador rojo si falla."""
    # Imagen embebida como data URI (base64)
    if src.startswith('data:'):
        try:
            import base64
            header, b64data = src.split(',', 1)
            datos = base64.b64decode(b64data)
        except Exception as e:
            print(f"  [!] No se pudo decodificar data URI: {e}")
            agregar_marcador_imagen(doc, '(data URI)', alt)
            return
    else:
        img_url = urljoin(page_url, src)
        datos = download_image(session, img_url)
    if datos:
        try:
            doc.add_picture(io.BytesIO(datos), width=Inches(5.5))
            ultimo = doc.paragraphs[-1]
            ultimo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if alt:
                caption = doc.add_paragraph(alt)
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].italic = True
            return
        except Exception as e:
            print(f"  [!] No se pudo insertar imagen: {e}")
    agregar_marcador_imagen(doc, src, alt)


def agregar_contenido_al_doc(
    doc: Document,
    session: requests.Session,
    titulo: str,
    contenido,
    page_url: str,
):
    """Agrega el título y el contenido HTML al documento Word preservando el formato."""

    # Título de la sección
    heading = doc.add_heading(titulo, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if contenido is None:
        doc.add_paragraph("(No se encontró contenido en esta página)")
        doc.add_page_break()
        return

    def procesar_bloque(nodo):
        """Procesa un nodo de bloque y lo agrega al documento."""
        from bs4 import NavigableString

        nombre = getattr(nodo, 'name', None)

        # Nodo de texto suelto en un contenedor
        if isinstance(nodo, NavigableString):
            texto = str(nodo).strip()
            if texto:
                doc.add_paragraph(texto)
            return

        if not nombre or nombre in ('script', 'style', 'noscript', 'nav', 'footer', 'header', 'button'):
            return

        # ── Salto de línea ──
        if nombre == 'br':
            return

        # ── Encabezados ──
        if nombre in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            nivel = int(nombre[1])
            p = doc.add_heading('', level=min(nivel, 9))
            agregar_inline(nodo, p)
            return

        # ── Párrafo ──
        if nombre == 'p':
            # Si el párrafo contiene imágenes, procesarlas por separado
            hijos = list(nodo.children)
            tiene_img = any(getattr(h, 'name', None) == 'img' for h in hijos)
            if tiene_img:
                # Texto antes/después de la imagen y la imagen misma
                parrafo_actual = None
                for hijo in hijos:
                    if getattr(hijo, 'name', None) == 'img':
                        src = hijo.get('src', '') or hijo.get('data-src', '')
                        alt = hijo.get('alt', '').strip()
                        if src:
                            insertar_imagen(doc, session, src, alt, page_url)
                        parrafo_actual = None
                    else:
                        texto = str(hijo).strip() if isinstance(hijo, NavigableString) else hijo.get_text(strip=True)
                        if texto:
                            if parrafo_actual is None:
                                parrafo_actual = doc.add_paragraph()
                            agregar_inline(hijo, parrafo_actual)
            else:
                texto_total = nodo.get_text(strip=True)
                if texto_total:
                    p = doc.add_paragraph()
                    agregar_inline(nodo, p)
            return

        # ── Listas ──
        if nombre in ('ul', 'ol'):
            for li in nodo.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                agregar_inline(li, p)
            return

        # ── Imagen suelta ──
        if nombre == 'img':
            src = extraer_src_imagen(nodo)
            alt = nodo.get('alt', '').strip()
            if src:
                insertar_imagen(doc, session, src, alt, page_url)
            return

        # ── Picture (HTML5) ──
        if nombre == 'picture':
            img = nodo.find('img')
            if img:
                src = extraer_src_imagen(img)
                if not src:
                    source = nodo.find('source')
                    if source:
                        srcset = source.get('srcset', '')
                        src = srcset.split(',')[0].split()[0] if srcset else ''
                alt = img.get('alt', '').strip() if img else ''
                if src:
                    insertar_imagen(doc, session, src, alt, page_url)
            return

        # ── Video / iframe ──
        if nombre in ('iframe', 'video'):
            url_video = extraer_url_video(nodo)
            if url_video:
                titulo_video = nodo.get('title', '').strip()
                agregar_marcador_video(doc, url_video, titulo_video)
            return

        # ── Tabla ──
        if nombre == 'table':
            filas = nodo.find_all('tr')
            if not filas:
                return
            cols = max(len(f.find_all(['td', 'th'])) for f in filas)
            if cols == 0:
                return
            tabla = doc.add_table(rows=len(filas), cols=cols)
            tabla.style = 'Table Grid'
            for i, fila in enumerate(filas):
                celdas = fila.find_all(['td', 'th'])
                for j, celda in enumerate(celdas):
                    if j < cols:
                        celda_doc = tabla.rows[i].cells[j]
                        p = celda_doc.paragraphs[0]
                        agregar_inline(celda, p)
            doc.add_paragraph()
            return

        # ── Contenedor genérico (div, section, article, figure, span con img…) ──
        if tiene_hijos_bloque(nodo):
            for hijo in nodo.children:
                procesar_bloque(hijo)
        else:
            # Sin hijos de bloque: texto inline. Pero por si hay imgs muy anidadas
            # (p.ej. div > span > img) las buscamos como seguro
            imgs_profundas = nodo.find_all('img')
            for img in imgs_profundas:
                src = extraer_src_imagen(img)
                alt = img.get('alt', '').strip()
                if src:
                    insertar_imagen(doc, session, src, alt, page_url)
            # Texto del nodo (sin el texto de las imágenes, que es vacío de todas formas)
            if not imgs_profundas:
                texto_total = nodo.get_text(strip=True)
                if texto_total:
                    p = doc.add_paragraph()
                    agregar_inline(nodo, p)

    procesar_bloque(contenido)
    doc.add_page_break()


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  Scraper CENET — Genera .docx con contenido del curso")
    print("=" * 60)

    # Credenciales (sin mostrar la contraseña en pantalla)
    print("\nIngresá tus credenciales de cenet.inet.edu.ar")
    username = input("Usuario: ").strip()
    password = getpass.getpass("Contraseña: ")

    # URLs precargadas del curso: Lógica Cableada
    urls_precargadas = [
        "https://cenet.inet.edu.ar/mod/page/view.php?id=32693",  # 01. Introducción a la Lógica Cableada - Componentes
        "https://cenet.inet.edu.ar/mod/page/view.php?id=32694",  # 02. Ventajas y Desventajas de la Lógica Cableada
        "https://cenet.inet.edu.ar/mod/page/view.php?id=32695",  # 03. Relés Industriales y Contactores como Elementos de Control
        "https://cenet.inet.edu.ar/mod/page/view.php?id=32696",  # 04. Implementación de Funciones Lógicas
        "https://cenet.inet.edu.ar/mod/page/view.php?id=32697",  # 05. Retención o Memoria Eléctrica
        "https://cenet.inet.edu.ar/mod/page/view.php?id=32698",  # 06. Diseño de circuitos de control con Relés
        "https://cenet.inet.edu.ar/mod/page/view.php?id=32699",  # 07. Aplicación CAJAS ALTAS/BAJAS
    ]

    print("\nURLs precargadas del curso:")
    for i, u in enumerate(urls_precargadas, 1):
        print(f"  {i}. {u}")

    print("\n¿Querés agregar más URLs? (línea vacía para continuar con las de arriba)")
    urls = list(urls_precargadas)
    while True:
        url = input("URL adicional (o Enter para continuar): ").strip()
        if not url:
            break
        if url.startswith("http"):
            urls.append(url)
        else:
            print("  URL inválida, omitida.")

    # Carpeta de salida
    carpeta_salida = input("\nNombre de la carpeta donde guardar los archivos (Enter = carpeta actual): ").strip()
    if not carpeta_salida:
        carpeta_salida = os.path.dirname(os.path.abspath(__file__))
    else:
        carpeta_salida = os.path.join(os.path.dirname(os.path.abspath(__file__)), carpeta_salida)
        os.makedirs(carpeta_salida, exist_ok=True)

    # Crear sesión y loguearse
    session = requests.Session()
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/124.0.0.0 Safari/537.36"
    })

    print("\nIniciando sesión...")
    if not login(session, username, password):
        print("ERROR: Usuario o contraseña incorrectos. Verificá tus datos.")
        sys.exit(1)
    print("Sesión iniciada correctamente.")

    # Procesar cada URL — un .docx por página
    archivos_guardados = []
    for i, url in enumerate(urls, 1):
        print(f"\n[{i}/{len(urls)}] Descargando: {url}")
        try:
            titulo, contenido = get_page_content(session, url)
            print(f"  Título: {titulo}")

            # Crear documento individual
            doc = Document()
            estilo_normal = doc.styles["Normal"]
            estilo_normal.font.name = "Calibri"
            estilo_normal.font.size = Pt(11)

            agregar_contenido_al_doc(doc, session, titulo, contenido, url)

            # Nombre del archivo basado en el título de la página
            nombre_archivo = re.sub(r'[\\/*?:"<>|]', "_", titulo)
            nombre_archivo = nombre_archivo[:80]  # limitar longitud
            ruta_salida = os.path.join(carpeta_salida, f"{i:02d}_{nombre_archivo}.docx")

            doc.save(ruta_salida)
            archivos_guardados.append(ruta_salida)
            print(f"  Guardado: {os.path.basename(ruta_salida)}")

        except Exception as e:
            print(f"  ERROR al procesar la página: {e}")

    print(f"\n{'=' * 60}")
    print(f"  {len(archivos_guardados)} archivos guardados en:")
    print(f"  {carpeta_salida}")
    for ruta in archivos_guardados:
        print(f"    - {os.path.basename(ruta)}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
