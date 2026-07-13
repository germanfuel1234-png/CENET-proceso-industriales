"""  
generar_cuestionario.py
-----------------------
Descarga el contenido de las páginas 2 a 5 del curso CENET y genera
un cuestionario con 3 preguntas de opción múltiple por tema usando Groq (gratis).
Guarda el resultado en cuestionario_U1.docx

Dependencias:
    pip install requests beautifulsoup4 python-docx lxml openai

Clave Groq GRATIS en: https://console.groq.com  (registrate con Gmail)

Uso:
    python generar_cuestionario.py
"""

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openai import OpenAI
import getpass
import os
import re
import sys
import json
from urllib.parse import urljoin

BASE_URL  = "https://cenet.inet.edu.ar"
LOGIN_URL = f"{BASE_URL}/login/index.php"

# Páginas del curso (solo 2 a 5)
PAGINAS = [
    {"id": "32669", "titulo": "02. Estructura de un Sistema Automatizado"},
    {"id": "32670", "titulo": "03. Evolución de los Sistemas Automatizados"},
    {"id": "32671", "titulo": "04. Tipos de Sistema, Medios y Señales"},
    {"id": "32672", "titulo": "05. Sensores Industriales"},
]

PREGUNTAS_POR_TEMA = 3

# ════════════════════════════════════════════════════════════
# CONFIGURACIÓN — completá estos tres valores y listo
# ════════════════════════════════════════════════════════════
CENET_USUARIO   = "36157362"
CENET_PASSWORD  = "36157362"
GROQ_API_KEY    = "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"
# ════════════════════════════════════════════════════════════


# ─── Autenticación ────────────────────────────────────────────────────────────

def login(session: requests.Session, username: str, password: str) -> bool:
    resp = session.get(LOGIN_URL, timeout=30)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "lxml")
    token_input = soup.find("input", {"name": "logintoken"})
    logintoken = token_input["value"] if token_input else ""
    payload = {
        "username": username,
        "password": password,
        "logintoken": logintoken,
        "anchor": "",
    }
    resp = session.post(LOGIN_URL, data=payload, timeout=30)
    resp.raise_for_status()
    if "login" in resp.url or ("loginerrormessage" in resp.text and "login" in resp.text[:2000].lower()):
        return False
    return True


# ─── Extracción de texto ──────────────────────────────────────────────────────

def obtener_texto_pagina(session: requests.Session, page_id: str) -> str:
    """Descarga la página y extrae solo el texto del contenido de la clase."""
    url = f"{BASE_URL}/mod/page/view.php?id={page_id}"
    resp = session.get(url, timeout=30)
    resp.raise_for_status()
    soup = BeautifulSoup(resp.text, "lxml")

    contenido = (
        soup.find("div", class_=lambda c: c and all(
            cls in c for cls in ["box", "generalbox", "center", "clearfix"]))
        or soup.find("div", {"id": "region-main"})
        or soup.find("div", {"role": "main"})
        or soup.find("main")
    )

    if not contenido:
        return ""

    # Eliminar scripts y estilos antes de extraer texto
    for tag in contenido.find_all(["script", "style", "nav", "button"]):
        tag.decompose()

    texto = contenido.get_text(separator="\n", strip=True)

    # Limpiar líneas vacías múltiples
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    return texto[:6000]  # limitar tokens (~4500 palabras es suficiente)


# ─── Generación de preguntas con OpenAI ───────────────────────────────────────

PROMPT_SISTEMA = """Eres un profesor especialista en automatización industrial.
Tu tarea es generar preguntas en español para evaluar la comprensión
del material de un curso técnico. Sé claro, preciso y usa terminología técnica apropiada."""

PROMPT_MULTIPLE = """Basándote EXCLUSIVAMENTE en el siguiente texto del tema "{titulo}",
genera exactamente {n} preguntas de opción múltiple en español.

Reglas:
- Cada pregunta debe tener 4 opciones (A, B, C, D)
- Solo una opción es correcta
- Las opciones incorrectas deben ser plausibles pero claramente erróneas para quien leyó el material
- No hagas preguntas triviales ni de cultura general

Devuelve ÚNICAMENTE un JSON válido con esta estructura exacta (sin texto extra):
[
  {{
    "pregunta": "Texto de la pregunta",
    "opciones": {{
      "A": "Opción A",
      "B": "Opción B",
      "C": "Opción C",
      "D": "Opción D"
    }},
    "respuesta_correcta": "A"
  }}
]

TEXTO DEL MATERIAL:
{texto}
"""

PROMPT_INVESTIGACION = """Basándote en el siguiente texto del tema "{titulo}",
genera exactamente {n} preguntas de investigación en español.

Las preguntas de investigación deben:
- Requerir que el alumno busque información más allá del texto dado
- Invitar a profundizar, comparar, relacionar o aplicar conceptos
- Ser preguntas abiertas, sin respuesta única (no sí/no)
- Estar redactadas con verbos como: investigar, analizar, comparar, explicar, describir, relacionar, proponer
- Tener una orientación práctica o de aplicación industrial cuando sea posible

Devuelve ÚNICAMENTE un JSON válido con esta estructura (sin texto extra):
[
  {{
    "pregunta": "Texto de la pregunta de investigación"
  }}
]

TEXTO DEL MATERIAL:
{texto}
"""


def generar_preguntas(client: OpenAI, titulo: str, texto: str, n: int) -> list:
    """Genera preguntas de opción múltiple con Groq."""
    prompt = PROMPT_MULTIPLE.format(titulo=titulo, n=n, texto=texto)

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",  # modelo gratuito de Groq
        messages=[
            {"role": "system", "content": PROMPT_SISTEMA},
            {"role": "user",   "content": prompt},
        ],
        temperature=0.7,
        max_tokens=2000,
    )

    raw = response.choices[0].message.content.strip()

    # Extraer el JSON aunque venga envuelto en markdown ```json ... ```
    m = re.search(r'\[.*\]', raw, re.DOTALL)
    if m:
        raw = m.group(0)

    return json.loads(raw)


INV_POR_TEMA = 2  # preguntas de investigación por tema


def generar_preguntas_investigacion(client: OpenAI, titulo: str, texto: str, n: int) -> list:
    """Genera preguntas de investigación abiertas con Groq."""
    prompt = PROMPT_INVESTIGACION.format(titulo=titulo, n=n, texto=texto)

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": PROMPT_SISTEMA},
            {"role": "user",   "content": prompt},
        ],
        temperature=0.8,
        max_tokens=1000,
    )

    raw = response.choices[0].message.content.strip()
    m = re.search(r'\[.*\]', raw, re.DOTALL)
    if m:
        raw = m.group(0)
    return json.loads(raw)


INV_POR_TEMA = 2  # preguntas de investigación por tema

def construir_docx(cuestionario: list, ruta: str):
    """
    cuestionario: lista de dicts con keys: titulo, preguntas
    Genera un .docx con preguntas y un anexo de respuestas al final.
    """
    doc = Document()

    # ── Estilos ──
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # ── Portada ──
    titulo_doc = doc.add_heading("Cuestionario – Unidad 1", level=0)
    titulo_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Introducción a la Automatización Industrial")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    doc.add_paragraph()

    respuestas_finales = []  # para el anexo

    num_global = 1  # numeración continua de preguntas

    for bloque in cuestionario:
        # Título del tema
        h = doc.add_heading(bloque["titulo"], level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        # ── Opción múltiple ──
        subtitulo_mc = doc.add_heading("Parte A — Opción Múltiple", level=2)

        for pq in bloque["preguntas"]:
            # Enunciado
            p_enunciado = doc.add_paragraph()
            run_num = p_enunciado.add_run(f"{num_global}. ")
            run_num.bold = True
            run_num.font.size = Pt(11)
            run_texto = p_enunciado.add_run(pq["pregunta"])
            run_texto.bold = True
            run_texto.font.size = Pt(11)

            # Opciones
            letras = ["A", "B", "C", "D"]
            for letra in letras:
                opcion_texto = pq["opciones"].get(letra, "")
                if opcion_texto:
                    p_op = doc.add_paragraph(style="List Bullet")
                    run_letra = p_op.add_run(f"{letra})  ")
                    run_letra.bold = True
                    p_op.add_run(opcion_texto)

            doc.add_paragraph()

            respuestas_finales.append({
                "num": num_global,
                "tema": bloque["titulo"],
                "correcta": pq["respuesta_correcta"],
                "texto_correcto": pq["opciones"].get(pq["respuesta_correcta"], ""),
            })
            num_global += 1

        # ── Preguntas de investigación ──
        if bloque.get("investigacion"):
            doc.add_paragraph()
            subtitulo_inv = doc.add_heading("Parte B — Preguntas de Investigación", level=2)
            for idx, pq_inv in enumerate(bloque["investigacion"], 1):
                p_inv = doc.add_paragraph()
                run_n = p_inv.add_run(f"{idx}. ")
                run_n.bold = True
                run_n.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                run_q = p_inv.add_run(pq_inv["pregunta"])
                run_q.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
                # Espacio para respuesta
                for _ in range(4):
                    p_linea = doc.add_paragraph("_" * 90)
                    p_linea.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
                doc.add_paragraph()

        doc.add_page_break()

    # ── Anexo: Respuestas correctas ──
    doc.add_page_break()
    h_resp = doc.add_heading("Respuestas Correctas", level=1)
    h_resp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Table Grid"

    # Encabezado de tabla
    encabezados = ["N°", "Respuesta", "Opción correcta"]
    for i, enc in enumerate(encabezados):
        cell = tabla.rows[0].cells[i]
        cell.text = enc
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Fondo azul oscuro
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E79')
        tc_pr.append(shd)

    for r in respuestas_finales:
        fila = tabla.add_row().cells
        fila[0].text = str(r["num"])
        fila[1].text = r["correcta"]
        fila[2].text = r["texto_correcto"]
        fila[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        fila[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(ruta)


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  Generador de Cuestionario CENET — U1 (temas 02 al 05)")
    print("=" * 60)

    # Leer credenciales de la sección de configuración
    username = CENET_USUARIO
    password = CENET_PASSWORD
    api_key  = GROQ_API_KEY

    if not api_key:
        print("\nERROR: falta la API Key de Groq.")
        print("Registrate gratis en https://console.groq.com (usá tu Gmail)")
        print("Después pegá la clave en GROQ_API_KEY dentro del archivo.")
        sys.exit(1)

    # Sesión CENET
    session = requests.Session()
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) "
                      "Chrome/124.0.0.0 Safari/537.36"
    })

    print("\nIniciando sesión en CENET...")
    if not login(session, username, password):
        print("ERROR: Usuario o contraseña incorrectos.")
        sys.exit(1)
    print("Sesión iniciada.")

    # Cliente Groq (compatible con la API de OpenAI)
    client = OpenAI(
        api_key=api_key,
        base_url="https://api.groq.com/openai/v1"
    )

    # Procesar cada página
    cuestionario = []
    for pagina in PAGINAS:
        print(f"\nProcesando: {pagina['titulo']}")
        print("  Descargando contenido...")
        texto = obtener_texto_pagina(session, pagina["id"])
        if not texto:
            print("  Sin contenido, omitida.")
            continue

        print(f"  Generando {PREGUNTAS_POR_TEMA} preguntas múltiple choice + {INV_POR_TEMA} de investigación con IA...")
        try:
            preguntas = generar_preguntas(client, pagina["titulo"], texto, PREGUNTAS_POR_TEMA)
            investigacion = generar_preguntas_investigacion(client, pagina["titulo"], texto, INV_POR_TEMA)
            cuestionario.append({
                "titulo": pagina["titulo"],
                "preguntas": preguntas,
                "investigacion": investigacion,
            })
            print(f"  {len(preguntas)} múltiple choice + {len(investigacion)} investigación generadas.")
        except Exception as e:
            print(f"  ERROR al generar preguntas: {e}")

    if not cuestionario:
        print("\nNo se pudo generar ningún cuestionario.")
        sys.exit(1)

    # Guardar .docx
    ruta_salida = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "cuestionario_U1.docx"
    )
    construir_docx(cuestionario, ruta_salida)

    total_preguntas = sum(len(b["preguntas"]) for b in cuestionario)
    print(f"\n{'=' * 60}")
    print(f"  Cuestionario generado: {total_preguntas} preguntas en total")
    print(f"  Guardado en: {ruta_salida}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
